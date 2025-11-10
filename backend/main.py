from __future__ import annotations

import io
import os
import secrets
from typing import Dict, Optional

import pandas as pd
from fastapi import Depends, FastAPI, File, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response, FileResponse

import schemas
from schemas import (
    ErrorResponse,
    MetricsRequest,
    MetricsResponse,
    TimeSeriesPoint,
    UploadResponse,
    FunnelRequest,
    FunnelResponse,
    FunnelPoint,
    StatTestRequest,
    StatTestResult,
    CohortAggregationResponse,
    CohortAggregationRow,
    CaptainLevelRequest,
    CaptainLevelResponse,
    CaptainLevelAggregationRow,
    MobileNumberUploadResponse,
    CaptainIdRequest,
    CaptainIdResponse,
    AOFunnelRequest,
    AOFunnelResponse,
    DaprBucketRequest,
    DaprBucketResponse,
    Fe2NetRequest,
    Fe2NetResponse,
    RtuPerformanceRequest,
    RtuPerformanceResponse,
    R2ARequest,
    R2AResponse,
    R2APercentageRequest,
    R2APercentageResponse,
    A2PhhSummaryRequest,
    A2PhhSummaryResponse,
    ReportItem,
    ReportAddRequest,
    ReportAddResponse,
    ReportUpdateCommentRequest,
    ReportUpdateTitleRequest,
    ReportListResponse,
    ReportExportResponse,
)
from transformations import (
    aggregate_time_series,
    filter_by_date_range,
    normalized_growth,
    rolling_average,
    subset_by_cohorts,
    get_cohort,
    compute_cohort_funnel_timeseries,
    compute_metric_timeseries_by_cohort,
)


app = FastAPI(title="Cohort Metrics API")

# Simple in-memory session store mapping session_id to DataFrame
SESSION_STORE: Dict[str, pd.DataFrame] = {}

# Funnel analysis session store - separate from main session store
FUNNEL_SESSION_STORE: Dict[str, pd.DataFrame] = {}

# Report builder session store - maps report_id to list of report items
REPORT_STORE: Dict[str, list[dict]] = {}

app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://internal-tools-v1-1.onrender.com",  # Your frontend URL
                    "http://localhost:5173", 
                    "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def get_session_df(x_session_id: Optional[str] = Header(default=None)) -> pd.DataFrame:
    if not x_session_id or x_session_id not in SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing session_id. Upload data first.")
    return SESSION_STORE[x_session_id]


@app.post("/upload", response_model=UploadResponse, responses={400: {"model": ErrorResponse}})
async def upload_csv(file: UploadFile = File(...)) -> UploadResponse:
    if not file.filename.lower().endswith(".csv"):
        raise HTTPException(status_code=400, detail="Only CSV files are supported")
    content = await file.read()
    try:
        df = pd.read_csv(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=400, detail=f"Failed to read CSV: {exc}")

    # Validate identifiers
    if "cohort" not in df.columns:
        raise HTTPException(status_code=400, detail="CSV missing required column: 'cohort'")
    if ("date" not in df.columns) and ("time" not in df.columns):
        raise HTTPException(status_code=400, detail="CSV must include 'date' (YYYY-MM-DD) or 'time' (YYYYMMDD)")

    df = df.copy()
    # Coerce a unified date column
    # df['date'] = df['time'].astype(str)
    if "date" in df.columns:
        df["date"] = pd.to_datetime(df["date"], errors="coerce")
    else:
        df["date"] = pd.to_datetime(df["time"].astype(str), format="%Y%m%d", errors="coerce")
    if df["date"].isna().any():
        x = df["date"].isna().sum()
        raise HTTPException(status_code=400, detail=f"Invalid date/time values found = {x}")
    df["cohort"] = df["cohort"].astype(str)
    # Best-effort numeric coercion for other columns
    for c in df.columns:
        if c in {"cohort", "date", "time"}:
            continue
        try:
            df[c] = pd.to_numeric(df[c], errors="ignore")
        except Exception:
            pass

    session_id = secrets.token_hex(16)
    SESSION_STORE[session_id] = df

    cohorts = sorted(df["cohort"].dropna().unique().tolist())
    date_min = df["date"].min()
    date_max = df["date"].max()

    # Determine available metric columns
    metric_candidates = [c for c in df.columns if c not in {"cohort", "date", "time"}]
    return UploadResponse(
        session_id=session_id,
        num_rows=df.shape[0],
        columns=list(df.columns.astype(str)),
        cohorts=cohorts,
        date_min=date_min.strftime("%Y-%m-%d"),
        date_max=date_max.strftime("%Y-%m-%d"),
        metrics=metric_candidates,
    )


@app.get("/meta")
def get_meta(df: pd.DataFrame = Depends(get_session_df)):
    cohorts = sorted(df["cohort"].dropna().unique().tolist())
    date_min = df["date"].min().strftime("%Y-%m-%d")
    date_max = df["date"].max().strftime("%Y-%m-%d")
    metrics = [c for c in df.columns if c not in {"cohort", "date", "time"}]
    
    # Identify categorical columns (non-numeric columns excluding cohort and date)
    categorical_columns = []
    for col in df.columns:
        if col not in {"cohort", "date", "time"}:
            # Check if column is categorical (object type or has limited unique values)
            if df[col].dtype == 'object' or df[col].dtype.name == 'category':
                categorical_columns.append(col)
            elif df[col].dtype in ['int64', 'int32', 'float64', 'float32']:
                # If numeric but has limited unique values (less than 20), consider it categorical
                unique_count = df[col].nunique()
                if unique_count < 20 and unique_count < len(df) * 0.1:
                    categorical_columns.append(col)
    
    return {
        "cohorts": cohorts,
        "date_min": date_min,
        "date_max": date_max,
        "metrics": metrics,
        "categorical_columns": sorted(categorical_columns)
    }


@app.post("/metrics", response_model=MetricsResponse, responses={400: {"model": ErrorResponse}})
def compute_metrics(payload: MetricsRequest, df: pd.DataFrame = Depends(get_session_df)) -> MetricsResponse:
    working = df.copy()

    # Apply cohort subset if provided
    cohorts = [c for c in [payload.test_cohort, payload.control_cohort] if c]
    if cohorts:
        working = subset_by_cohorts(working, cohorts)

    # Pre and Post period filters are used to compute summary aggregations per period
    pre_df = working
    post_df = working
    if payload.pre_period:
        pre_df = filter_by_date_range(working, payload.pre_period.start_date, payload.pre_period.end_date)
    if payload.post_period:
        post_df = filter_by_date_range(working, payload.post_period.start_date, payload.post_period.end_date)

    # Time series for plotting (full range filtered only by cohorts)
    ts_df = working.copy()
    ts_df = ts_df.sort_values(["cohort", "date"]).reset_index(drop=True)

    # Rolling windows
    for w in payload.rolling_windows:
        ts_df = rolling_average(ts_df, value_col="metric_value", window=w, by=["cohort"], date_col="date")

    # Normalized growth relative to baseline date or first date
    ts_df = normalized_growth(
        ts_df,
        value_col="metric_value",
        baseline_date=payload.normalized_growth_baseline_date,
        by=["cohort"],
        date_col="date",
    )

    # Build timeseries response points
    def _row_to_point(row) -> TimeSeriesPoint:
        point = {
            "date": pd.to_datetime(row["date"]).strftime("%Y-%m-%d"),
            "cohort": row["cohort"],
            "metric_value": float(row["metric_value"]),
            "metric_value_pct_change": float(row.get("metric_value_pct_change", float("nan"))) if "metric_value_pct_change" in row else None,
        }
        for w in payload.rolling_windows:
            col = f"metric_value_roll_{w}"
            if col in ts_df.columns:
                val = row.get(col)
                point[col] = float(val) if pd.notna(val) else None
        return TimeSeriesPoint(**point)

    time_series = [_row_to_point(r) for r in ts_df.to_dict("records")]

    # Compute summaries for pre and post periods by cohort
    summaries = []
    def _compute_agg(d: pd.DataFrame, label: str):
        for agg in payload.aggregations:
            out = aggregate_time_series(d, group_by=["cohort"], value_col="metric_value", agg=agg)
            # Map to test/control values
            test_val = out.loc[out["cohort"] == payload.test_cohort, f"metric_value_{agg}"]
            control_val = out.loc[out["cohort"] == payload.control_cohort, f"metric_value_{agg}"]
            tv = float(test_val.iloc[0]) if not test_val.empty else 0.0
            cv = float(control_val.iloc[0]) if not control_val.empty else 0.0
            mean_diff = tv - cv
            pct_change = (mean_diff / cv * 100.0) if cv != 0 else None
            summaries.append({
                "aggregation": agg,
                "test_value": tv,
                "control_value": cv,
                "mean_difference": mean_diff,
                "pct_change": pct_change,
            })

    _compute_agg(pre_df, "pre")
    _compute_agg(post_df, "post")

    return MetricsResponse(
        time_series=time_series,
        summaries=summaries,
    )


@app.post("/funnel", response_model=FunnelResponse, responses={400: {"model": ErrorResponse}})
def funnel(payload: FunnelRequest, df: pd.DataFrame = Depends(get_session_df)) -> FunnelResponse:
    working = df.copy()
    
    # Apply confirmation filtering if specified (support per-test/control and legacy)
    test_confirmed = getattr(payload, 'test_confirmed', None) or getattr(payload, 'confirmed', None) or ''
    control_confirmed = getattr(payload, 'control_confirmed', None) or getattr(payload, 'confirmed', None) or ''
    
    # Filter for test and control cohorts with optional confirmation filtering
    if payload.test_cohort and payload.control_cohort:
        test_data = get_cohort(working, payload.test_cohort, test_confirmed).copy()
        control_data = get_cohort(working, payload.control_cohort, control_confirmed).copy()
        # Label cohorts explicitly so identical cohort names remain distinct
        test_data["cohort"] = test_data["cohort"].astype(str).apply(lambda c: f"TEST: {c}")
        control_data["cohort"] = control_data["cohort"].astype(str).apply(lambda c: f"CONTROL: {c}")
        working = pd.concat([test_data, control_data], ignore_index=True)
    elif payload.test_cohort:
        working = get_cohort(working, payload.test_cohort, test_confirmed).copy()
        working["cohort"] = working["cohort"].astype(str).apply(lambda c: f"TEST: {c}")
    elif payload.control_cohort:
        working = get_cohort(working, payload.control_cohort, control_confirmed).copy()
        working["cohort"] = working["cohort"].astype(str).apply(lambda c: f"CONTROL: {c}")
    else:
        # If no specific cohorts, just apply confirmation filter if specified
        if test_confirmed:
            if test_confirmed not in working.columns:
                raise HTTPException(status_code=400, detail=f"Confirmation column '{test_confirmed}' not found")
            working = working[~working[test_confirmed].isna()]

    # Handle series breakout if specified
    series_breakout_col = payload.series_breakout
    if series_breakout_col:
        if series_breakout_col not in working.columns:
            raise HTTPException(status_code=400, detail=f"Series breakout column '{series_breakout_col}' not found in dataset. Available columns: {list(working.columns)[:20]}")
        # Check if column has any non-null values
        if working[series_breakout_col].isna().all():
            raise HTTPException(status_code=400, detail=f"Series breakout column '{series_breakout_col}' has no valid (non-null) values")
        # Store the series breakout column before computing timeseries
        series_values = working[series_breakout_col].astype(str).unique().tolist()
    else:
        series_values = None

    ts = compute_cohort_funnel_timeseries(working)
    metrics_available = [c for c in ts.columns if c not in {"date", "cohort"}]
    if not metrics_available:
        raise HTTPException(status_code=400, detail="No metrics available in dataset")

    # Determine the metric to use
    metric = payload.metric or metrics_available[0]
    agg = payload.agg or "sum"
    
    # If series breakout is specified, we need to recompute the metric with groupby
    # This is because precomputed metrics don't include the series breakout column
    if series_breakout_col:
        try:
            # Check if metric exists in working dataframe as a column
            if metric in working.columns:
                # Metric is a direct column, compute with groupby
                ts_with_series = compute_metric_timeseries_by_cohort(working, metric, agg, group_by=["cohort", series_breakout_col])
            else:
                # Metric is a precomputed one (like ao_days, online_days, etc.)
                # We need to compute it per series_breakout group from the original working data
                # by calling compute_cohort_funnel_timeseries on each group
                
                # Ensure date column exists in working
                working_for_series = working.copy()
                if "date" not in working_for_series.columns:
                    if "time" in working_for_series.columns:
                        working_for_series["date"] = pd.to_datetime(working_for_series["time"].astype(str), format="%Y%m%d", errors="coerce")
                    else:
                        raise HTTPException(status_code=400, detail="Cannot compute series breakout: missing 'date' or 'time' column")
                
                # Group working by series_breakout and compute timeseries for each group
                frames = []
                series_vals = working_for_series[series_breakout_col].dropna().unique()
                if len(series_vals) == 0:
                    raise HTTPException(status_code=400, detail=f"Series breakout column '{series_breakout_col}' has no valid values")
                
                for series_val in series_vals:
                    group_df = working_for_series[working_for_series[series_breakout_col] == series_val].copy()
                    if len(group_df) == 0:
                        continue
                    try:
                        group_ts = compute_cohort_funnel_timeseries(group_df)
                        if metric in group_ts.columns:
                            group_ts[series_breakout_col] = str(series_val)
                            frames.append(group_ts[["date", "cohort", metric, series_breakout_col]])
                    except Exception as group_error:
                        # Log but continue with other groups
                        print(f"Warning: Failed to compute timeseries for series value '{series_val}': {group_error}")
                        continue
                
                if frames:
                    ts_with_series = pd.concat(frames, ignore_index=True)
                else:
                    raise HTTPException(status_code=400, detail=f"Cannot compute series breakout for metric '{metric}': metric not found in any grouped data")
            
            # Replace ts with ts_with_series (which includes series_breakout)
            ts = ts_with_series.copy()
            
        except HTTPException:
            raise
        except Exception as e:
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=400, detail=f"Failed to compute series breakout: {str(e)}")
    elif payload.metric and payload.metric not in metrics_available:
        # No series breakout, but metric needs to be computed
        try:
            extra = compute_metric_timeseries_by_cohort(working, payload.metric, agg)
        except Exception as e:
            raise HTTPException(status_code=400, detail=str(e))
        # Merge with existing to reuse date filtering & response shape
        ts = ts.copy()
        merged = ts.merge(extra, on=["date", "cohort"], how="outer", suffixes=("", "_extra"))
        # If both existed, prefer the non-null new computation
        if f"{metric}_extra" in merged.columns and metric in merged.columns:
            merged[metric] = merged[metric].fillna(merged[f"{metric}_extra"]).astype(float)
            merged = merged.drop(columns=[f"{metric}_extra"]) 
        elif f"{metric}_extra" in merged.columns and metric not in merged.columns:
            merged = merged.rename(columns={f"{metric}_extra": metric})
        ts = merged
        metrics_available = [c for c in ts.columns if c not in {"date", "cohort"}]

    pre_df = ts
    post_df = ts
    if payload.pre_period:
        pre_df = filter_by_date_range(ts, payload.pre_period.start_date, payload.pre_period.end_date, date_col="date")
    if payload.post_period:
        post_df = filter_by_date_range(ts, payload.post_period.start_date, payload.post_period.end_date, date_col="date")

    def to_points(d: pd.DataFrame) -> list[FunnelPoint]:
        points = []
        for r in d.to_dict("records"):
            series_val = str(r.get(series_breakout_col, "")) if series_breakout_col else None
            points.append(FunnelPoint(
                date=pd.to_datetime(r["date"]).strftime("%Y-%m-%d"),
                cohort=r["cohort"],
                metric=metric,
                value=float(r.get(metric, 0.0)),
                series_value=series_val
            ))
        return points

    def summarize(d: pd.DataFrame) -> dict[str, float]:
        if d.empty:
            return {}
        tmp = d.groupby("cohort")[metric].sum().reset_index()
        return {str(row["cohort"]): float(row[metric]) for _, row in tmp.iterrows()}

    return FunnelResponse(
        metrics_available=metrics_available,
        pre_series=to_points(pre_df),
        post_series=to_points(post_df),
        pre_summary=summarize(pre_df),
        post_summary=summarize(post_df),
    )


@app.get("/cohort-aggregation", response_model=CohortAggregationResponse, responses={400: {"model": ErrorResponse}})
def get_cohort_aggregation(df: pd.DataFrame = Depends(get_session_df)) -> CohortAggregationResponse:
    """Get cohort-level aggregation data as shown in the example"""
    working = df.copy()

    # Check if required columns exist
    required_columns = [
        "totalExpCaps", "visitedCaps", "clickedCaptain", "count_captain_pitch_centre_card_clicked_city","count_captain_pitch_centre_card_visible_city", "exploredCaptains",
        "exploredCaptains_Subs", "exploredCaptains_EPKM", "exploredCaptains_FlatCommission",
        "exploredCaptains_CM", "confirmedCaptains", "confirmedCaptains_Subs",
        "confirmedCaptains_Subs_purchased", "confirmedCaptains_Subs_purchased_weekend",
        "confirmedCaptains_EPKM", "confirmedCaptains_FlatCommission", "confirmedCaptains_CM"
    ]

    missing_columns = [col for col in required_columns if col not in working.columns]
    if missing_columns:
        raise HTTPException(
            status_code=400,
            detail=f"Missing required columns for cohort aggregation: {missing_columns}. "
                   f"Available columns: {list(working.columns)}"
        )

    # Perform the aggregation as specified in the example
    result = working.groupby(["cohort"]).agg({
        "totalExpCaps": "nunique",
        "visitedCaps": "nunique",
        'clickedCaptain': 'nunique',
        'count_captain_pitch_centre_card_clicked_city': 'sum',
        'count_captain_pitch_centre_card_visible_city': 'sum',
        'exploredCaptains': 'nunique',
        'exploredCaptains_Subs': 'nunique',
        'exploredCaptains_EPKM': 'nunique',
        'exploredCaptains_FlatCommission': 'nunique',
        'exploredCaptains_CM': 'nunique',
        'confirmedCaptains': 'nunique',
        'confirmedCaptains_Subs': 'nunique',
        'confirmedCaptains_Subs_purchased': 'nunique',
        'confirmedCaptains_Subs_purchased_weekend': 'nunique',
        'confirmedCaptains_EPKM': 'nunique',
        'confirmedCaptains_FlatCommission': 'nunique',
        'confirmedCaptains_CM': 'nunique'

    }).reset_index().sort_values("exploredCaptains", ascending=False)

    # Calculate the ratio columns
    result['Visit2Click'] = result['clickedCaptain'] / result['visitedCaps']
    result['Base2Visit'] = result['visitedCaps'] / result['totalExpCaps']
    # To fix division by zero, use np.where to safely compute Click2Confirm
    import numpy as np
    result['Click2Confirm'] = np.where(result['clickedCaptain'] == 0, 0, result['confirmedCaptains'] / result['clickedCaptain'])

    # Handle division by zero
    result['Visit2Click'] = result['Visit2Click'].fillna(0)
    result['Base2Visit'] = result['Base2Visit'].fillna(0)
    result['Click2Confirm'] = result['Click2Confirm'].fillna(0)
    # Convert to list of CohortAggregationRow objects
    data = []
    for _, row in result.iterrows():
        data.append(CohortAggregationRow(
            cohort=str(row['cohort']),
            totalExpCaps=float(row['totalExpCaps']),
            visitedCaps=float(row['visitedCaps']),
            clickedCaptain=float(row['clickedCaptain']),
            pitch_centre_card_clicked=float(row['count_captain_pitch_centre_card_clicked_city']),
            pitch_centre_card_visible=float(row['count_captain_pitch_centre_card_visible_city']),
            exploredCaptains=float(row['exploredCaptains']),
            exploredCaptains_Subs=float(row['exploredCaptains_Subs']),
            exploredCaptains_EPKM=float(row['exploredCaptains_EPKM']),
            exploredCaptains_FlatCommission=float(row['exploredCaptains_FlatCommission']),
            exploredCaptains_CM=float(row['exploredCaptains_CM']),
            confirmedCaptains=float(row['confirmedCaptains']),
            confirmedCaptains_Subs=float(row['confirmedCaptains_Subs']),
            confirmedCaptains_Subs_purchased=float(row['confirmedCaptains_Subs_purchased']),
            confirmedCaptains_Subs_purchased_weekend=float(row['confirmedCaptains_Subs_purchased_weekend']),
            confirmedCaptains_EPKM=float(row['confirmedCaptains_EPKM']),
            confirmedCaptains_FlatCommission=float(row['confirmedCaptains_FlatCommission']),
            confirmedCaptains_CM=float(row['confirmedCaptains_CM']),
            Visit2Click=float(row['Visit2Click']),
            Base2Visit=float(row['Base2Visit']),
            Click2Confirm=float(row['Click2Confirm']),
        ))

    return CohortAggregationResponse(data=data)


@app.delete("/session")
def clear_session(x_session_id: Optional[str] = Header(default=None)):
    if x_session_id and x_session_id in SESSION_STORE:
        del SESSION_STORE[x_session_id]
    return {"ok": True}


@app.post("/statistical-test")
def run_statistical_test_endpoint(
    request: StatTestRequest,
    x_session_id: Optional[str] = Header(default=None)
) -> StatTestResult:
    """Run statistical tests on cohort data"""
    from statistical_analysis import run_statistical_test
    
    try:
        result = run_statistical_test(request)
        return result
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/captain-level-aggregation", response_model=CaptainLevelResponse, responses={400: {"model": ErrorResponse}})
def captain_level_aggregation(
    payload: CaptainLevelRequest,
    df: pd.DataFrame = Depends(get_session_df)
) -> CaptainLevelResponse:
    """
    Captain-level aggregation grouped by time and a categorical column.
    Filters data by cohort and confirmation level, then aggregates metrics.
    """
    working = df.copy()
    
    # Ensure date column is present
    if "date" not in working.columns:
        raise HTTPException(status_code=400, detail="Dataset must have 'date' column")
    
    # Ensure group_by_column exists
    if payload.group_by_column not in working.columns:
        raise HTTPException(
            status_code=400,
            detail=f"Group by column '{payload.group_by_column}' not found in dataset"
        )
    
    # Validate all metric columns exist
    for metric_agg in payload.metric_aggregations:
        if metric_agg.column not in working.columns:
            raise HTTPException(
                status_code=400,
                detail=f"Metric column '{metric_agg.column}' not found in dataset"
            )
    
    def filter_captain_level(cohort_name: str, confirmation_filter: Optional[str]) -> pd.DataFrame:
        """Filter data by cohort and optional confirmation level"""
        cohort_df = working[working["cohort"] == cohort_name].copy()
        if confirmation_filter and confirmation_filter in cohort_df.columns:
            cohort_df = cohort_df[~cohort_df[confirmation_filter].isna()]
        return cohort_df
    
    # Get test and control data
    test_df = filter_captain_level(payload.test_cohort, payload.test_confirmed)
    control_df = filter_captain_level(payload.control_cohort, payload.control_confirmed)
    
    if test_df.empty:
        raise HTTPException(status_code=400, detail=f"No data found for test cohort '{payload.test_cohort}'")
    if control_df.empty:
        raise HTTPException(status_code=400, detail=f"No data found for control cohort '{payload.control_cohort}'")
    
    # Helper to aggregate data
    def aggregate_data(data: pd.DataFrame, period: str, cohort_type: str) -> list[CaptainLevelAggregationRow]:
        """Group by date and group_by_column, then aggregate metrics"""
        if data.empty:
            return []
        
        # Build aggregation dict - map column names to aggregation functions
        agg_dict = {}
        agg_key_mapping = {}  # Maps original column to our custom key
        
        for metric_agg in payload.metric_aggregations:
            col_name = metric_agg.column
            agg_func = metric_agg.agg_func
            agg_key = f"{col_name}_{agg_func}"
            
            # For pandas groupby, we need column name as key
            if col_name not in agg_dict:
                agg_dict[col_name] = []
            agg_dict[col_name].append(agg_func)
            
            # Store mapping for later renaming
            agg_key_mapping[f"{col_name}_{agg_func}"] = agg_key
        
        # Group by date and the categorical column
        grouped = data.groupby(["date", payload.group_by_column]).agg(agg_dict).reset_index()
        
        # Flatten multi-level column names if they exist
        if isinstance(grouped.columns, pd.MultiIndex):
            new_cols = ["date", payload.group_by_column]
            for col in grouped.columns[2:]:  # Skip date and group_by_column
                if col[1]:  # If there's an aggregation function
                    new_cols.append(f"{col[0]}_{col[1]}")
                else:
                    new_cols.append(col[0])
            grouped.columns = new_cols
        
        # Convert to response format
        rows = []
        for _, row in grouped.iterrows():
            aggregations = {}
            for orig_key, custom_key in agg_key_mapping.items():
                val = row.get(orig_key)
                aggregations[custom_key] = float(val) if pd.notna(val) else 0.0
            
            rows.append(CaptainLevelAggregationRow(
                period=period,
                cohort_type=cohort_type,
                date=pd.to_datetime(row["date"]).strftime("%Y-%m-%d"),
                group_value=str(row[payload.group_by_column]),
                aggregations=aggregations
            ))
        
        return rows
    
    # Filter by date ranges
    pre_test = test_df
    post_test = test_df
    pre_control = control_df
    post_control = control_df
    
    if payload.pre_period:
        pre_test = filter_by_date_range(
            test_df,
            payload.pre_period.start_date,
            payload.pre_period.end_date,
            date_col="date"
        )
        pre_control = filter_by_date_range(
            control_df,
            payload.pre_period.start_date,
            payload.pre_period.end_date,
            date_col="date"
        )
    
    if payload.post_period:
        post_test = filter_by_date_range(
            test_df,
            payload.post_period.start_date,
            payload.post_period.end_date,
            date_col="date"
        )
        post_control = filter_by_date_range(
            control_df,
            payload.post_period.start_date,
            payload.post_period.end_date,
            date_col="date"
        )
    
    # Aggregate all combinations
    result_data = []
    result_data.extend(aggregate_data(pre_test, "pre", "test"))
    result_data.extend(aggregate_data(post_test, "post", "test"))
    result_data.extend(aggregate_data(pre_control, "pre", "control"))
    result_data.extend(aggregate_data(post_control, "post", "control"))
    
    # Extract metric names
    metrics = [f"{m.column}_{m.agg_func}" for m in payload.metric_aggregations]
    
    return CaptainLevelResponse(
        data=result_data,
        group_by_column=payload.group_by_column,
        metrics=metrics
    )


# ============================================================================
# FUNNEL ANALYSIS ENDPOINTS
# ============================================================================

@app.post("/funnel-analysis/upload-mobile-numbers", response_model=MobileNumberUploadResponse, responses={400: {"model": ErrorResponse}})
async def upload_mobile_numbers(file: UploadFile = File(...)) -> MobileNumberUploadResponse:
    """
    Upload CSV with mobile_number column (and optional cohort column) for funnel analysis
    """
    if not file.filename.lower().endswith(".csv"):
        raise HTTPException(status_code=400, detail="Only CSV files are supported")
    
    content = await file.read()
    try:
        df = pd.read_csv(io.BytesIO(content))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to read CSV: {exc}")
    
    # Validate that mobile_number column exists
    if "mobile_number" not in df.columns:
        raise HTTPException(status_code=400, detail="CSV must include 'mobile_number' column")
    
    # Check if cohort column exists
    has_cohort = "cohort" in df.columns
    
    # Ensure mobile_number is treated as string
    df['mobile_number'] = df['mobile_number'].astype(str)
    
    if has_cohort:
        df['cohort'] = df['cohort'].astype(str)
    
    # Drop duplicates based on mobile_number (and cohort if present)
    original_rows = len(df)
    if has_cohort:
        df = df.drop_duplicates(subset=['mobile_number', 'cohort'], keep='first')
    else:
        df = df.drop_duplicates(subset=['mobile_number'], keep='first')
    
    duplicates_removed = original_rows - len(df)
    
    # Generate session ID and store
    funnel_session_id = secrets.token_hex(16)
    FUNNEL_SESSION_STORE[funnel_session_id] = df
    
    # Get preview (first 5 rows)
    preview = df.head(5).to_dict('records')
    
    return MobileNumberUploadResponse(
        funnel_session_id=funnel_session_id,
        num_rows=len(df),
        columns=list(df.columns),
        has_cohort=has_cohort,
        preview=preview,
        duplicates_removed=duplicates_removed
    )


@app.post("/funnel-analysis/get-captain-ids", response_model=CaptainIdResponse, responses={400: {"model": ErrorResponse}})
async def get_captain_ids(
    payload: CaptainIdRequest,
    x_funnel_session_id: Optional[str] = Header(default=None)
) -> CaptainIdResponse:
    """
    Fetch captain IDs for mobile numbers in current session
    """
    if not x_funnel_session_id or x_funnel_session_id not in FUNNEL_SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing funnel_session_id. Upload mobile numbers first.")
    
    mobile_number_df = FUNNEL_SESSION_STORE[x_funnel_session_id]
    
    try:
        from funnel import get_captain_id
        result_df = get_captain_id(mobile_number_df, payload.username)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch captain IDs from Presto: {exc}")
    
    # Update session with new dataframe including captain_id
    FUNNEL_SESSION_STORE[x_funnel_session_id] = result_df
    
    # Count how many captain IDs were found (non-null)
    num_captains_found = len(result_df['captain_id'].dropna().unique())
    
    # Get preview
    preview = result_df.head(5).to_dict('records')
    
    return CaptainIdResponse(
        num_rows=len(result_df),
        num_captains_found=int(num_captains_found),
        preview=preview
    )


@app.post("/funnel-analysis/get-ao-funnel", response_model=AOFunnelResponse, responses={400: {"model": ErrorResponse}})
async def get_ao_funnel_data(
    payload: AOFunnelRequest,
    x_funnel_session_id: Optional[str] = Header(default=None)
) -> AOFunnelResponse:
    """
    Fetch AO funnel metrics for captain IDs in current session
    """
    if not x_funnel_session_id or x_funnel_session_id not in FUNNEL_SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing funnel_session_id. Get captain IDs first.")
    
    captain_id_df = FUNNEL_SESSION_STORE[x_funnel_session_id]
    
    # Validate that captain_id column exists
    if 'captain_id' not in captain_id_df.columns:
        raise HTTPException(status_code=400, detail="No captain_id column found. Run 'Get Captain IDs' first.")
    
    # Check if there are any valid captain IDs
    if captain_id_df['captain_id'].isna().all():
        raise HTTPException(status_code=400, detail="No valid captain IDs found. Cannot fetch funnel data.")
    
    try:
        from funnel import get_ao_funnel
        result_df = get_ao_funnel(
            captain_id_df,
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.time_level,
            payload.tod_level
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch AO funnel data from Presto: {exc}")
    
    # Format the dataframe for cohort analysis
    # Ensure we have a proper date column
    if 'time' in result_df.columns and 'date' not in result_df.columns:
        # Convert time column (YYYYMMDD format) to date
        result_df['date'] = pd.to_datetime(result_df['time'], format='%Y%m%d', errors='coerce')
    elif 'time' in result_df.columns:
        # Also ensure date column is datetime
        result_df['date'] = pd.to_datetime(result_df['date'], errors='coerce')
    
    # Ensure cohort column exists, if not create from existing data
    if 'cohort' not in result_df.columns:
        # If no cohort column, create a default one
        result_df['cohort'] = 'all_captains'
    
    # Update session with funnel data
    FUNNEL_SESSION_STORE[x_funnel_session_id] = result_df
    
    # Identify metric columns (exclude identifier columns)
    exclude_cols = {'mobile_number', 'captain_id', 'cohort', 'city', 'time', 'date'}
    metric_cols = [c for c in result_df.columns if c not in exclude_cols]
    
    # Calculate unique captain IDs from full dataset
    unique_captain_ids = int(result_df['captain_id'].nunique())
    
    # Get preview (first 10 rows)
    preview = result_df.head(10).to_dict('records')
    
    return AOFunnelResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        preview=preview,
        metrics=metric_cols,
        unique_captain_ids=unique_captain_ids
    )


@app.post("/funnel-analysis/use-for-analysis", response_model=UploadResponse, responses={400: {"model": ErrorResponse}})
def use_funnel_for_analysis(x_funnel_session_id: Optional[str] = Header(default=None)) -> UploadResponse:
    """
    Transfer funnel analysis data to main cohort analysis session
    This allows using the AO funnel data for plotting and further analysis
    """
    if not x_funnel_session_id or x_funnel_session_id not in FUNNEL_SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing funnel_session_id. Complete funnel analysis first.")
    
    # Get the funnel dataframe
    funnel_df = FUNNEL_SESSION_STORE[x_funnel_session_id].copy()
    
    # Ensure date column is properly formatted
    if 'date' not in funnel_df.columns and 'time' in funnel_df.columns:
        funnel_df['date'] = pd.to_datetime(funnel_df['time'], format='%Y%m%d', errors='coerce')
    elif 'date' in funnel_df.columns:
        funnel_df['date'] = pd.to_datetime(funnel_df['date'], errors='coerce')
    else:
        raise HTTPException(status_code=400, detail="No date/time column found in funnel data")
    
    # Ensure cohort column exists
    if 'cohort' not in funnel_df.columns:
        funnel_df['cohort'] = 'all_captains'
    
    # Ensure cohort is string
    funnel_df['cohort'] = funnel_df['cohort'].astype(str)
    
    # Drop any rows with invalid dates
    invalid_dates = funnel_df['date'].isna().sum()
    if invalid_dates > 0:
        funnel_df = funnel_df[~funnel_df['date'].isna()]
    
    # Create a new session in the main store
    session_id = secrets.token_hex(16)
    SESSION_STORE[session_id] = funnel_df
    
    # Get cohorts and date range
    cohorts = sorted(funnel_df["cohort"].dropna().unique().tolist())
    date_min = funnel_df["date"].min()
    date_max = funnel_df["date"].max()
    
    # Determine available metric columns
    exclude_cols = {'cohort', 'date', 'time', 'mobile_number', 'captain_id', 'city'}
    metric_candidates = [c for c in funnel_df.columns if c not in exclude_cols]
    
    return UploadResponse(
        session_id=session_id,
        num_rows=funnel_df.shape[0],
        columns=list(funnel_df.columns.astype(str)),
        cohorts=cohorts,
        date_min=date_min.strftime("%Y-%m-%d"),
        date_max=date_max.strftime("%Y-%m-%d"),
        metrics=metric_candidates,
    )


@app.get("/funnel-analysis/export-csv")
def export_funnel_csv(x_funnel_session_id: Optional[str] = Header(default=None)):
    """Export full funnel dataset as CSV"""
    from fastapi.responses import StreamingResponse
    
    if not x_funnel_session_id or x_funnel_session_id not in FUNNEL_SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing funnel_session_id")
    
    df = FUNNEL_SESSION_STORE[x_funnel_session_id]
    
    # Convert to CSV
    csv_buffer = io.StringIO()
    df.to_csv(csv_buffer, index=False)
    csv_buffer.seek(0)
    
    return StreamingResponse(
        iter([csv_buffer.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=funnel_data.csv"}
    )


@app.post("/funnel-analysis/dapr-bucket", response_model=DaprBucketResponse, responses={400: {"model": ErrorResponse}})
async def get_dapr_bucket(payload: DaprBucketRequest) -> DaprBucketResponse:
    """
    Fetch DAPR bucket distribution data from Presto
    """
    try:
        from funnel import dapr_bucket
        result_df = dapr_bucket(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.service_category,
            payload.low_dapr,
            payload.high_dapr
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch DAPR bucket data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return DaprBucketResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.post("/captain-dashboards/fe2net", response_model=Fe2NetResponse, responses={400: {"model": ErrorResponse}})
async def get_fe2net(payload: Fe2NetRequest) -> Fe2NetResponse:
    """
    Fetch FE2Net funnel data from Presto
    """
    try:
        from funnel import fe2net
        result_df = fe2net(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.service_category,
            payload.geo_level,
            payload.time_level
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch FE2Net data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return Fe2NetResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.post("/captain-dashboards/rtu-performance", response_model=RtuPerformanceResponse, responses={400: {"model": ErrorResponse}})
async def get_rtu_performance(payload: RtuPerformanceRequest) -> RtuPerformanceResponse:
    """
    Fetch RTU Performance metrics from Presto
    """
    try:
        from funnel import performance_metrics
        result_df = performance_metrics(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.perf_cut,
            payload.consistency_cut,
            payload.time_level,
            payload.tod_level,
            payload.service_category
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch RTU Performance data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return RtuPerformanceResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.post("/captain-dashboards/r2a", response_model=R2AResponse, responses={400: {"model": ErrorResponse}})
async def get_r2a(payload: R2ARequest) -> R2AResponse:
    """
    Fetch R2A% (Registration to Activation) metrics from Presto
    """
    try:
        from funnel import r2a_registration_by_activation
        result_df = r2a_registration_by_activation(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.service,
            payload.time_level
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch R2A data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return R2AResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.post("/captain-dashboards/r2a-percentage", response_model=R2APercentageResponse, responses={400: {"model": ErrorResponse}})
async def get_r2a_percentage(payload: R2APercentageRequest) -> R2APercentageResponse:
    """
    Fetch R2A% metrics from Presto
    """
    try:
        from funnel import r2a_pecentage
        result_df = r2a_pecentage(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.service,
            payload.time_level
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch R2A% data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return R2APercentageResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.post("/captain-dashboards/a2phh-summary", response_model=A2PhhSummaryResponse, responses={400: {"model": ErrorResponse}})
async def get_a2phh_summary(payload: A2PhhSummaryRequest) -> A2PhhSummaryResponse:
    """
    Fetch A2PHH Summary M0 metrics from Presto
    """
    try:
        from funnel import a2phh_summary
        result_df = a2phh_summary(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.service,
            payload.time_level
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch A2PHH Summary data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return A2PhhSummaryResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.delete("/funnel-analysis/session")
def clear_funnel_session(x_funnel_session_id: Optional[str] = Header(default=None)):
    """Clear funnel analysis session"""
    if x_funnel_session_id and x_funnel_session_id in FUNNEL_SESSION_STORE:
        del FUNNEL_SESSION_STORE[x_funnel_session_id]
    return {"ok": True}


# ============================================================================
# REPORT BUILDER ENDPOINTS
# ============================================================================

@app.post("/report/create")
def create_report() -> dict:
    """Create a new report session"""
    report_id = secrets.token_hex(16)
    REPORT_STORE[report_id] = []
    return {"report_id": report_id}


@app.post("/report/add-item", response_model=ReportAddResponse)
def add_report_item(
    payload: ReportAddRequest,
    x_report_id: Optional[str] = Header(default=None, alias="x-report-id")
) -> ReportAddResponse:
    """Add an item (chart, table, or text) to the report"""
    from datetime import datetime
    
    print(f"DEBUG: add_report_item called with x_report_id={x_report_id}")
    print(f"DEBUG: Payload type={payload.type}, title={payload.title}")
    print(f"DEBUG: REPORT_STORE keys before: {list(REPORT_STORE.keys())}")
    
    # Create report if not exists
    if not x_report_id or x_report_id not in REPORT_STORE:
        if not x_report_id:
            x_report_id = secrets.token_hex(16)
        REPORT_STORE[x_report_id] = []
        print(f"DEBUG: Created new report with ID: {x_report_id}")
    
    item_id = secrets.token_hex(8)
    item = {
        "id": item_id,
        "type": payload.type,
        "title": payload.title,
        "content": payload.content,
        "comment": payload.comment or "",
        "timestamp": datetime.now().isoformat()
    }
    
    REPORT_STORE[x_report_id].append(item)
    print(f"DEBUG: Added item to report {x_report_id}, total items: {len(REPORT_STORE[x_report_id])}")
    print(f"DEBUG: REPORT_STORE keys after: {list(REPORT_STORE.keys())}")
    
    return ReportAddResponse(
        report_id=x_report_id,
        item_id=item_id,
        num_items=len(REPORT_STORE[x_report_id])
    )


@app.put("/report/update-comment")
def update_report_comment(
    payload: ReportUpdateCommentRequest,
    x_report_id: Optional[str] = Header(default=None, alias="x-report-id")
):
    """Update comment for a report item"""
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    items = REPORT_STORE[x_report_id]
    for item in items:
        if item["id"] == payload.item_id:
            item["comment"] = payload.comment
            return {"ok": True}
    
    raise HTTPException(status_code=404, detail="Item not found in report")


@app.put("/report/update-title")
def update_report_title(
    payload: ReportUpdateTitleRequest,
    x_report_id: Optional[str] = Header(default=None, alias="x-report-id")
):
    """Update title for a report item"""
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    items = REPORT_STORE[x_report_id]
    for item in items:
        if item["id"] == payload.item_id:
            item["title"] = payload.title
            return {"ok": True}
    
    raise HTTPException(status_code=404, detail="Item not found in report")


@app.delete("/report/item/{item_id}")
def delete_report_item(
    item_id: str,
    x_report_id: Optional[str] = Header(default=None, alias="x-report-id")
):
    """Delete an item from the report"""
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    items = REPORT_STORE[x_report_id]
    REPORT_STORE[x_report_id] = [item for item in items if item["id"] != item_id]
    
    return {"ok": True, "num_items": len(REPORT_STORE[x_report_id])}


@app.get("/report/list", response_model=ReportListResponse)
def list_report_items(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")) -> ReportListResponse:
    """Get all items in the current report"""
    print(f"DEBUG: list_report_items called with x_report_id={x_report_id}")
    print(f"DEBUG: REPORT_STORE keys: {list(REPORT_STORE.keys())}")
    
    if not x_report_id:
        print("DEBUG: No report ID provided, returning empty")
        return ReportListResponse(report_id="", items=[])
    
    if x_report_id not in REPORT_STORE:
        print(f"DEBUG: Report ID {x_report_id} not found in store, returning empty")
        return ReportListResponse(report_id=x_report_id, items=[])
    
    items = [ReportItem(**item) for item in REPORT_STORE[x_report_id]]
    print(f"DEBUG: Returning {len(items)} items for report {x_report_id}")
    return ReportListResponse(report_id=x_report_id, items=items)


@app.get("/report/export", response_model=ReportExportResponse)
def export_report(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")) -> ReportExportResponse:
    """Export report as HTML document"""
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    items = REPORT_STORE[x_report_id]
    
    # Build HTML document
    html_parts = [
        """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Experiment Report</title>
            <style>
                body {
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                    max-width: 1200px;
                    margin: 0 auto;
                    padding: 40px 20px;
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                }
                .container {
                    background: white;
                    padding: 40px;
                    border-radius: 12px;
                    box-shadow: 0 20px 60px rgba(0,0,0,0.3);
                }
                h1 {
                    color: #1e293b;
                    font-size: 32px;
                    margin-bottom: 10px;
                    border-bottom: 3px solid #8b5cf6;
                    padding-bottom: 10px;
                }
                .subtitle {
                    color: #64748b;
                    font-size: 14px;
                    margin-bottom: 30px;
                }
                .report-item {
                    margin: 30px 0;
                    padding: 25px;
                    border: 2px solid #e2e8f0;
                    border-radius: 8px;
                    background: #f8fafc;
                }
                .item-header {
                    font-size: 20px;
                    font-weight: 700;
                    color: #334155;
                    margin-bottom: 10px;
                }
                .item-type {
                    display: inline-block;
                    padding: 4px 12px;
                    background: #8b5cf6;
                    color: white;
                    border-radius: 12px;
                    font-size: 12px;
                    font-weight: 600;
                    margin-bottom: 12px;
                }
                .item-comment {
                    margin-top: 15px;
                    padding: 15px;
                    background: white;
                    border-left: 4px solid #8b5cf6;
                    font-style: italic;
                    color: #475569;
                }
                .item-timestamp {
                    font-size: 12px;
                    color: #94a3b8;
                    margin-top: 8px;
                }
                table {
                    width: 100%;
                    border-collapse: collapse;
                    margin-top: 15px;
                    font-size: 13px;
                }
                th {
                    background: #8b5cf6;
                    color: white;
                    padding: 10px;
                    text-align: left;
                    font-weight: 600;
                }
                td {
                    padding: 10px;
                    border-bottom: 1px solid #e2e8f0;
                }
                tr:nth-child(even) {
                    background: #f1f5f9;
                }
                .chart-config {
                    background: #f1f5f9;
                    padding: 15px;
                    border-radius: 6px;
                    margin-top: 15px;
                    font-family: 'Courier New', monospace;
                    font-size: 12px;
                }
                .page-break {
                    page-break-after: always;
                }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>📊 Experiment Report</h1>
                <p class="subtitle">Generated on """ + pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S") + """</p>
        """
    ]
    
    # Add each report item
    for idx, item in enumerate(items):
        item_html = f"""
            <div class="report-item">
                <span class="item-type">{item['type'].upper()}</span>
                <div class="item-header">{idx + 1}. {item['title']}</div>
        """
        
        # Render content based on type
        if item['type'] == 'table':
            # Render table image if available, otherwise render table data
            content = item['content']
            if 'imageDataUrl' in content and content['imageDataUrl']:
                # Display the actual table image
                img_data = content['imageDataUrl']
                item_html += f"<div style='margin-top:15px; text-align:center;'><img src='{img_data}' alt='{item['title']}' style='max-width:100%; height:auto; border:1px solid #e2e8f0; border-radius:8px;' /></div>"
            elif 'data' in content and len(content['data']) > 0:
                # Fallback to table data if no image
                # Get columns from first row
                columns = list(content['data'][0].keys())
                item_html += "<table>"
                item_html += "<thead><tr>"
                for col in columns:
                    item_html += f"<th>{col}</th>"
                item_html += "</tr></thead><tbody>"
                
                # Add rows (limit to first 50 for readability)
                for row in content['data'][:50]:
                    item_html += "<tr>"
                    for col in columns:
                        val = row.get(col, '')
                        # Format numbers
                        if isinstance(val, (int, float)):
                            val = f"{val:,.2f}" if isinstance(val, float) else f"{val:,}"
                        item_html += f"<td>{val}</td>"
                    item_html += "</tr>"
                item_html += "</tbody></table>"
                
                if len(content['data']) > 50:
                    item_html += f"<p style='margin-top:10px; color:#64748b; font-size:12px;'>Showing first 50 of {len(content['data'])} rows</p>"
        
        elif item['type'] == 'chart':
            # Render chart image if available
            content = item['content']
            if 'imageDataUrl' in content and content['imageDataUrl']:
                # Display the actual chart image
                img_data = content['imageDataUrl']
                item_html += f"<div style='margin-top:15px; text-align:center;'><img src='{img_data}' alt='{item['title']}' style='max-width:100%; height:auto; border:1px solid #e2e8f0; border-radius:8px;' /></div>"
            else:
                # Fallback to metadata if no image
                item_html += "<div class='chart-config'>"
                item_html += f"<strong>Chart Type:</strong> {content.get('chartType', 'N/A')}<br>"
                item_html += f"<strong>X-Axis:</strong> {content.get('xAxis', 'N/A')}<br>"
                item_html += f"<strong>Y-Axes:</strong> {', '.join(content.get('yAxes', []))}<br>"
                if content.get('seriesBy'):
                    item_html += f"<strong>Series Breakout:</strong> {content.get('seriesBy')}<br>"
                item_html += f"<strong>Data Points:</strong> {len(content.get('data', []))}"
                item_html += "</div>"
        
        elif item['type'] == 'text':
            # Render text content
            text_content = item['content'].get('text', '')
            item_html += f"<div style='margin-top:15px; padding:15px; background:white; border-radius:6px;'>{text_content}</div>"
        
        # Add comment if present
        if item['comment']:
            item_html += f"<div class='item-comment'>💬 {item['comment']}</div>"
        
        # Add timestamp
        item_html += f"<div class='item-timestamp'>Added: {item['timestamp']}</div>"
        
        item_html += "</div>"
        
        # Add page break after every 2 items for printing
        if (idx + 1) % 2 == 0:
            item_html += "<div class='page-break'></div>"
        
        html_parts.append(item_html)
    
    html_parts.append("""
            </div>
        </body>
        </html>
    """)
    
    return ReportExportResponse(report_html="".join(html_parts))


@app.delete("/report/clear")
def clear_report(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")):
    """Clear all items from the report"""
    if x_report_id and x_report_id in REPORT_STORE:
        REPORT_STORE[x_report_id] = []
    return {"ok": True}


@app.get("/report/export/pdf")
def export_report_pdf(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")):
    """Export report as PDF"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
        from reportlab.lib.enums import TA_LEFT
        import base64
        import tempfile
        import os
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"PDF export library not available: {str(e)}")
    
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    try:
        items = REPORT_STORE[x_report_id]
        temp_files = []  # Track temporary files for cleanup
        
        # Create temporary PDF file
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Container for the 'Flowable' objects
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='black',
            spaceAfter=12,
        )
        
        # Add title
        story.append(Paragraph("Experiment Report", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Add each item
        for idx, item in enumerate(items):
            # Item title
            story.append(Paragraph(f"{idx + 1}. {item['title']}", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            # Add image if present
            if item['type'] in ['chart', 'table'] and 'imageDataUrl' in item.get('content', {}):
                try:
                    # Decode base64 image
                    img_data = item['content']['imageDataUrl']
                    if img_data.startswith('data:image'):
                        img_data = img_data.split(',')[1]
                    
                    img_bytes = base64.b64decode(img_data)
                    
                    # Use PIL to open the image, then save to a temporary file that stays open
                    from PIL import Image as PILImage
                    pil_img = PILImage.open(io.BytesIO(img_bytes))
                    
                    # Save to temporary file that ReportLab can read
                    img_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                    pil_img.save(img_file.name, format='PNG')
                    img_file.close()
                    
                    # Add image to PDF - ReportLab will read the file
                    img = RLImage(img_file.name, width=5*inch, height=3.75*inch)
                    story.append(img)
                    story.append(Spacer(1, 0.2*inch))
                    
                    # Clean up after PDF is built (will be done later)
                    # Store file path for cleanup
                    temp_files.append(img_file.name)
                except Exception as e:
                    print(f"Error adding image to PDF: {e}")
                    import traceback
                    traceback.print_exc()
            
            # Add text content
            if item['type'] == 'text' and 'text' in item.get('content', {}):
                story.append(Paragraph(item['content']['text'], styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
            
            # Add comment if present
            if item.get('comment'):
                comment_style = ParagraphStyle(
                    'Comment',
                    parent=styles['Normal'],
                    fontStyle='italic',
                    textColor='#666666',
                    leftIndent=20,
                    backColor='#FFFACD',
                )
                story.append(Paragraph(f"Comment: {item['comment']}", comment_style))
            
            story.append(Spacer(1, 0.3*inch))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        # Clean up temporary image files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except Exception as e:
                print(f"Error cleaning up temp file {temp_file}: {e}")
        
        return Response(
            content=buffer.read(),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=experiment_report_{pd.Timestamp.now().strftime('%Y%m%d')}.pdf",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "x-report-id",
            }
        )
    except Exception as e:
        # Clean up temporary files on error
        if 'temp_files' in locals():
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.unlink(temp_file)
                except:
                    pass
        
        print(f"Error generating PDF: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")


@app.get("/report/export/png")
def export_report_png(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")):
    """Export report as PNG (single page image)"""
    from PIL import Image, ImageDraw, ImageFont
    import base64
    import tempfile
    
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    items = REPORT_STORE[x_report_id]
    
    # Create a white background image
    width, height = 2480, 3508  # A4 size at 300 DPI
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    try:
        # Try to use a nice font
        try:
            font_title = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 60)
            font_text = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 40)
        except:
            font_title = ImageFont.load_default()
            font_text = ImageFont.load_default()
        
        y_offset = 100
        
        # Add title
        draw.text((100, y_offset), "Experiment Report", fill='black', font=font_title)
        y_offset += 150
        
        # Add each item
        for idx, item in enumerate(items):
            # Item title
            draw.text((100, y_offset), f"{idx + 1}. {item['title']}", fill='black', font=font_text)
            y_offset += 80
            
            # Add image if present
            if item['type'] in ['chart', 'table'] and 'imageDataUrl' in item.get('content', {}):
                try:
                    img_data = item['content']['imageDataUrl']
                    if img_data.startswith('data:image'):
                        img_data = img_data.split(',')[1]
                    
                    item_img_bytes = base64.b64decode(img_data)
                    item_img = Image.open(io.BytesIO(item_img_bytes))
                    
                    # Resize to fit
                    max_width = width - 200
                    max_height = 600
                    # Use Resampling.LANCZOS if available, otherwise fall back to LANCZOS
                    try:
                        item_img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                    except AttributeError:
                        item_img.thumbnail((max_width, max_height), Image.LANCZOS)
                    
                    # Paste image
                    img.paste(item_img, (100, y_offset))
                    y_offset += item_img.height + 50
                except Exception as e:
                    print(f"Error adding image: {e}")
            
            # Add text
            if item['type'] == 'text' and 'text' in item.get('content', {}):
                text = item['content']['text'][:200]  # Limit text length
                draw.text((100, y_offset), text, fill='black', font=font_text)
                y_offset += 100
            
            # Add comment if present
            if item.get('comment'):
                comment_y = y_offset
                # Calculate comment box height based on text length
                comment_text = f"Comment: {item['comment'][:150]}"
                # Estimate height: ~40 pixels per line, max 3 lines
                lines = comment_text.split('\n')
                num_lines = min(len(lines), 3)
                comment_box_height = max(60, num_lines * 40)
                
                # Draw comment box background (yellow highlight)
                draw.rectangle([90, comment_y - 5, width - 100, comment_y + comment_box_height], fill='#FFFACD', outline='#FFD700', width=2)
                
                # Draw comment text (wrap if needed)
                try:
                    # Try to use textbbox if available (Pillow >= 8.0)
                    if hasattr(draw, 'textbbox'):
                        bbox = draw.textbbox((100, comment_y), comment_text, font=font_text)
                        actual_height = bbox[3] - bbox[1] + 20
                        comment_box_height = max(comment_box_height, actual_height)
                except:
                    pass
                
                # Draw comment text
                draw.text((100, comment_y), comment_text, fill='#333333', font=font_text)
                y_offset += comment_box_height + 20
            
            y_offset += 100
            
            if y_offset > height - 200:
                break  # Stop if we run out of space
        
    except Exception as e:
        print(f"Error creating PNG: {e}")
    
    # Save to bytes
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    buffer.seek(0)
    
    return Response(
        content=buffer.read(),
        media_type="image/png",
        headers={"Content-Disposition": f"attachment; filename=experiment_report_{pd.Timestamp.now().strftime('%Y%m%d')}.png"}
    )


@app.get("/report/export/word")
def export_report_word(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")):
    """Export report as Word document"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import base64
        import tempfile
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"Word export library not available: {str(e)}")
    
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    try:
        items = REPORT_STORE[x_report_id]
        
        # Create document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add title
        title = doc.add_heading('Experiment Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add each item
        for idx, item in enumerate(items):
            # Item title
            doc.add_heading(f"{idx + 1}. {item['title']}", level=1)
            
            # Add image if present
            if item['type'] in ['chart', 'table'] and 'imageDataUrl' in item.get('content', {}):
                try:
                    img_data = item['content']['imageDataUrl']
                    if img_data.startswith('data:image'):
                        img_data = img_data.split(',')[1]
                    
                    img_bytes = base64.b64decode(img_data)
                    img_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                    img_file.write(img_bytes)
                    img_file.close()
                    
                    # Add image to document
                    doc.add_picture(img_file.name, width=Inches(6))
                    
                    # Clean up
                    os.unlink(img_file.name)
                except Exception as e:
                    print(f"Error adding image to Word: {e}")
            
            # Add text content
            if item['type'] == 'text' and 'text' in item.get('content', {}):
                doc.add_paragraph(item['content']['text'])
            
            # Add comment if present
            if item.get('comment'):
                para = doc.add_paragraph(f"Comment: {item['comment']}")
                para.italic = True
            
            # Add spacing
            doc.add_paragraph()
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return Response(
            content=buffer.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=experiment_report_{pd.Timestamp.now().strftime('%Y%m%d')}.docx",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "x-report-id",
            }
        )
    except Exception as e:
        print(f"Error generating Word document: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(e)}")


@app.get("/health")
def health():
    return {"status": "ok"}


if __name__ == "__main__":
    import uvicorn
    # Render sets PORT environment variable, default to 8000 for local development
    port = int(os.environ.get("PORT", "8000"))
    uvicorn.run(app, host="0.0.0.0", port=port)

